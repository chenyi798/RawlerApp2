import requests
from urllib.parse import urlencode
import time
import random
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
import os
from urllib.parse import urljoin
import re

class PBCCrawler:
    def __init__(self, log_callback=None):
        self.log_callback = log_callback if log_callback else print
        self.is_crawling = True
        
        # 配置
        self.config = {
            'min_delay': 2,
            'max_delay': 4,
            'max_retries': 3,
            'timeout': 30,
            'max_pages': 10000000  # 限制最大页数
        }
    
    def log(self, message, level="INFO"):
        """日志记录"""
        self.log_callback(f"{message}", level)
    
    def stop_crawling(self):
        """停止爬取"""
        self.is_crawling = False
        self.log("收到停止信号", "WARNING")
    
    def crawl_pbc_search(self, page=1, keyword="金融监管", retry_count=0):
        """
        爬取中国人民银行网站搜索结果的函数
        """
        if not self.is_crawling:
            return None
            
        url = "http://wzdig.pbc.gov.cn:8080/search/pcRender"
        
        params = {
            "pageId": "fa445f64514c40c68b1c8ffe859c649e"
        }
        
        data = {
            "originalSearchUrl": "/search/pcRender?pageId=fa445f64514c40c68b1c8ffe859c649e",
            "originalSearch": "",
            "app": "9fc3239692e448f894e3bd8d674b55b8,fea427d898234869be673fce4767b655,c103cb417c3e4ca79b7660f11e19cc8a,8d97b205b3844c58a6fad8f846106c39,d774a7e1668d4c7ebbb2ec96519ab466,07ea3d4cf62a4953aacbed6ff295c37a,f4c26e783ca448fe815ee3353fc9ab54,da38e4920e16490d91643876517fc623,a02114069e134838847c9e23090d8a41,1b78b68b865c4da58372d1a9f04b9782,4063d07b489d4a9a874170f97fb7600c,f5390b519b7144bea54dbb7078f3545f,910172d1ac5a4a1f8122291df54b1cf2,71a08f51290d488782be1575a4c56b05,14fc626387524a019fb092604b88d2db,d445ce2a46144608831a7beff59010f0,bba920ef834447c78f3e235429db9375,73094b4d50c14db984741455360531c0,2487d9db44e24c98bf0ca496a8ccc525,3c6575ffd059462c982b187810a2ab42,c9831f1dad3a4fd5a7e7ee1b152b0744,2167d3331f4045098ebdffddb9f04cab,e1ea1c866c82452a95b1b2a8547733a6,21424b94fd544684891bb477f813fe38,a576e41580e942ecb77736150058f2f3,13f3bfaf44a3438bb555c8a740995c4e,161f5d5b817d409bb5cbd4c5aebe48b0,079971477b754c5eb91a05087e244623,43da8aa849cc41d5a098c458dd4741d9,b1463b021ebe4531b57712e876f9d7fc,de28fe0e9b4a4ffc93958fe74072cb08,f29e12211177456a9ed0657e6e5c19c3,64ffe8bb793c445bb07a9036326a69eb,bcaf7d31fe164a2ea6d75fff601a008e,b9e0c76b7bf24d0da12ab4755893184c,9641e32d9d2c4410a397dd85463f6923,d74612e724a342a5b4cc8d4793342d48,7a5471bbd80a4a0cbbebd02555714f2a,88c9d6458fcd4d9baf16cfc1f0782532,2d342bb6f743463d88be3c85c1e9fae2,506771dc122b4d80aaa3ef93bec39e61,50436d23f0954c48932b50687f271984,847a6664a5424143949f9dbbda5c0b9d",
            "appName": "",
            "sr": "score desc",
            "advtime": "",
            "advrange": "",
            "ext": "-siteId:3688005",
            "pNo": str(page),
            "advepq": "",
            "advoq": "",
            "adveq": "",
            "searchArea": "",
            "advSiteArea": "",
            "q": keyword
        }
        
        headers = {
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.7",
            "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8,en-GB;q=0.7,en-US;q=0.6",
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "Content-Type": "application/x-www-form-urlencoded",
            "Cookie": "JSESSIONID=0000CNLQ0SY8wmNRiZ6FSDUyPbw:-1",
            "Host": "wzdig.pbc.gov.cn:8080",
            "Origin": "http://wzdig.pbc.gov.cn:8080",
            "Pragma": "no-cache",
            "Upgrade-Insecure-Requests": "1",
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/142.0.0.0 Safari/537.36 Edg/142.0.0.0"
        }
        
        try:
            response = requests.post(
                url=url,
                params=params,
                data=data,
                headers=headers,
                timeout=self.config['timeout']
            )
            
            if response.status_code == 200:
                self.log(f"第{page}页爬取成功！")
                return response.text
            else:
                self.log(f"请求失败，状态码：{response.status_code}", "WARNING")
                return None
                
        except requests.exceptions.RequestException as e:
            self.log(f"请求发生错误：{e}", "WARNING")
            
            # 重试机制
            if retry_count < self.config['max_retries']:
                retry_delay = random.uniform(2, 5)
                self.log(f"第{retry_count+1}次重试，等待{retry_delay:.1f}秒...")
                time.sleep(retry_delay)
                return self.crawl_pbc_search(page, keyword, retry_count+1)
            else:
                self.log(f"已达到最大重试次数({self.config['max_retries']})，放弃请求", "ERROR")
                return None

    def get_total_pages(self, html_content):
        """
        从搜索结果页面获取总页数
        """
        if not html_content:
            return 1
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # 查找包含总页数的元素
        total_records_span = soup.find('span', class_='default-result-tolal-records')
        
        if total_records_span:
            # 提取内部的span文本内容
            inner_span = total_records_span.find('span')
            if inner_span:
                try:
                    total_pages = int(inner_span.get_text().strip())
                    self.log(f"检测到总页数: {total_pages}")
                    return total_pages
                except ValueError:
                    self.log(f"无法解析页数: {inner_span.get_text()}", "WARNING")
        
        # 如果无法获取总页数，返回默认值1
        self.log("无法获取总页数，使用默认值1", "WARNING")
        return 1

    def extract_links_with_titles_from_result_list(self, html_content):
        """
        提取class="default-result-list conMid_con"下面的a链接和标题
        """
        if not html_content:
            return []
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        result_list_div = soup.find('div', class_='default-result-list conMid_con')
        
        if not result_list_div:
            self.log("未找到class='default-result-list conMid_con'的div", "WARNING")
            return []
        
        a_tags = result_list_div.find_all('a', href=True)
        
        link_data = []
        for a_tag in a_tags:
            href = a_tag.get('href', '').strip()
            title = a_tag.get_text(strip=True)
            
            if href and title:
                if href.startswith('http://www.pbc.gov.cnhttps://'):
                    href = href.replace('http://www.pbc.gov.cnhttps://', 'https://')
                
                link_data.append({
                    'title': title,
                    'link': href
                })
        
        return link_data

    def extract_title(self, html_content):
        """
        从HTML中提取标题
        """
        try:
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # 多种选择器尝试
            selectors = [
                'h2[style*="font-size: 16px"][style*="color: #333"]',
                'h1',
                'h2',
                '.title',
                '.article-title',
                'title'
            ]
            
            for selector in selectors:
                title_tag = soup.select_one(selector)
                if title_tag:
                    title_text = title_tag.get_text().strip()
                    if title_text and len(title_text) > 5:  # 避免过短的标题
                        if selector != selectors[0]:
                            self.log(f"使用备用选择器找到标题: {selector}")
                        return title_text
            
            return "未知标题"
            
        except Exception as e:
            self.log(f"提取标题时出错: {e}", "WARNING")
            return "未知标题"

    def extract_content(self, html_content):
        """
        从HTML内容中提取正文内容
        优先提取id="UCAP-CONTENT"的div，如果没有则提取id="zoom"的div
        """
        try:
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # 首先尝试提取id="UCAP-CONTENT"的div
            ucap_content_div = soup.find('div', id='UCAP-CONTENT')
            
            if ucap_content_div:
                self.log("使用id='UCAP-CONTENT'的内容")
                return {
                    'html': ucap_content_div.prettify(),
                    'text': ucap_content_div.get_text(separator='\n', strip=True)
                }
            
            # 如果没有id="UCAP-CONTENT"，则尝试提取id="zoom"的div
            zoom_div = soup.find('div', id='zoom')
            
            if zoom_div:
                self.log("使用id='zoom'的内容")
                return {
                    'html': zoom_div.prettify(),
                    'text': zoom_div.get_text(separator='\n', strip=True)
                }
            
            # 尝试其他常见的内容选择器
            content_selectors = [
                '.content',
                '.article-content',
                '.main-content',
                '.text-content',
                'div[class*="content"]',
                'div[class*="article"]'
            ]
            
            for selector in content_selectors:
                content_div = soup.select_one(selector)
                if content_div:
                    self.log(f"使用选择器找到内容: {selector}")
                    return {
                        'html': content_div.prettify(),
                        'text': content_div.get_text(separator='\n', strip=True)
                    }
            
            # 如果两者都没有找到
            self.log("未找到特定内容区域，将使用body内容", "WARNING")
            body_content = soup.find('body')
            if body_content:
                return {
                    'html': body_content.prettify(),
                    'text': body_content.get_text(separator='\n', strip=True)
                }
            
            return {'html': '', 'text': ''}
        
        except Exception as e:
            self.log(f"提取内容时出错: {e}", "WARNING")
            return {'html': '', 'text': ''}

    def download_file_with_retry(self, file_url, base_url=None, max_retries=2, timeout=10):
        """
        下载文件，支持重试机制
        """
        if not self.is_crawling:
            return None
            
        # 如果是相对路径，组合成完整URL
        if base_url and not file_url.startswith(('http://', 'https://')):
            file_url = urljoin(base_url, file_url)
        
        for attempt in range(max_retries + 1):
            if not self.is_crawling:
                return None
                
            try:
                headers = {
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/142.0.0.0 Safari/537.36 Edg/142.0.0.0',
                    'Referer': base_url or 'http://www.pbc.gov.cn/'
                }
                
                response = requests.get(file_url, headers=headers, timeout=timeout)
                if response.status_code == 200:
                    return response.content
                else:
                    self.log(f"文件下载失败，状态码：{response.status_code}，URL：{file_url}", "WARNING")
            except requests.exceptions.Timeout:
                self.log(f"文件下载超时（尝试 {attempt+1}/{max_retries+1}），URL：{file_url}", "WARNING")
            except requests.exceptions.RequestException as e:
                self.log(f"文件下载错误（尝试 {attempt+1}/{max_retries+1}）：{e}，URL：{file_url}", "WARNING")
            
            # 如果不是最后一次尝试，等待一段时间后重试
            if attempt < max_retries:
                time.sleep(1)
        
        return None

    def get_article_base_url(self, article_url):
        """
        从文章URL提取基础URL（目录部分）
        """
        # 移除文件名部分，保留目录
        if '/' in article_url:
            # 找到最后一个斜杠的位置
            last_slash_index = article_url.rfind('/')
            if last_slash_index >= 0:
                return article_url[:last_slash_index + 1]
        return article_url

    def clean_filename(self, filename, max_length=100):
        """
        清理文件名，移除非法字符
        """
        if not filename:
            return "untitled"
        
        # 移除HTML标签
        clean_name = re.sub(r'<[^>]+>', '', filename)
        
        # 移除Windows文件名中不允许的字符
        illegal_chars = r'[<>:"/\\|?*]'
        clean_name = re.sub(illegal_chars, '', clean_name)
        
        # 移除多余的空格和换行符
        clean_name = re.sub(r'\s+', ' ', clean_name).strip()
        
        # 限制文件名长度
        if len(clean_name) > max_length:
            clean_name = clean_name[:max_length].rstrip()
        
        # 如果清理后为空，使用默认名称
        if not clean_name:
            clean_name = "untitled"
        
        return clean_name

    def save_html_to_doc(self, html_content, doc_filename, article_url):
        """
        将HTML内容保存到Word文档，有图片就下载保存，没有就直接保存文本
        """
        if not self.is_crawling:
            return False
            
        try:
            doc = Document()
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # 从文章URL获取基础URL（目录部分）
            article_base_url = self.get_article_base_url(article_url)
            
            # 检查是否有图片
            img_elements = soup.find_all('img')
            total_images = len(img_elements)
            
            if total_images > 0:
                # 有图片的情况：分别处理文本和图片
                self.log(f"发现 {total_images} 张图片，将下载并保存到文档")
                
                # 统计图片下载情况
                successful_images = 0
                
                # 先处理所有文本内容
                all_text = soup.get_text(separator='\n', strip=True)
                if all_text:
                    lines = all_text.split('\n')
                    for line in lines:
                        if line.strip():
                            doc.add_paragraph(line.strip())
                
                # 然后处理图片
                for i, img_element in enumerate(img_elements, 1):
                    if not self.is_crawling:
                        break
                        
                    img_src = img_element.get('src')
                    if img_src:
                        self.log(f"正在下载第 {i}/{total_images} 张图片: {img_src}")
                        # 使用文章基础URL作为图片下载的基础URL
                        image_content = self.download_file_with_retry(img_src, article_base_url, max_retries=2, timeout=10)
                        
                        if image_content:
                            try:
                                # 保存图片到临时文件
                                temp_img_path = f'temp_image_{i}.jpg'
                                with open(temp_img_path, 'wb') as f:
                                    f.write(image_content)
                                
                                # 添加到Word文档
                                doc.add_picture(temp_img_path, width=Inches(6))
                                successful_images += 1
                                
                                # 删除临时文件
                                os.remove(temp_img_path)
                                
                            except Exception as e:
                                self.log(f"处理图片文件失败 {img_src}: {e}", "WARNING")
                                doc.add_paragraph(f"[图片加载失败: {img_src}]")
                        else:
                            doc.add_paragraph(f"[图片下载失败: {img_src}]")
                
                # 保存文档
                doc.save(doc_filename)
                self.log(f"文档已保存为: {doc_filename}，图片下载成功率: {successful_images}/{total_images}")
                return True
                
            else:
                # 没有图片的情况：直接保存文本内容
                self.log("没有发现图片，直接保存文本内容")
                all_text = soup.get_text(separator='\n', strip=True)
                
                if all_text:
                    lines = all_text.split('\n')
                    for line in lines:
                        if line.strip():
                            doc.add_paragraph(line.strip())
                    
                    doc.save(doc_filename)
                    self.log(f"文档已保存为: {doc_filename}，共 {len(lines)} 行文本")
                    return True
                else:
                    self.log("警告：没有提取到任何文本内容", "WARNING")
                    # 保存一个空的文档或包含提示信息的文档
                    doc.add_paragraph("该页面没有提取到任何文本内容")
                    doc.save(doc_filename)
                    return True
        
        except Exception as e:
            self.log(f"保存文档时出错: {e}", "ERROR")
            return False

    def download_excel_files(self, html_content, article_url, output_folder, title):
        """
        从HTML内容中提取并下载Excel文件
        """
        if not self.is_crawling:
            return 0
            
        try:
            soup = BeautifulSoup(html_content, 'html.parser')
            article_base_url = self.get_article_base_url(article_url)
            
            # 查找所有链接，筛选出Excel文件
            excel_extensions = ['.xls', '.xlsx', '.csv']
            all_links = soup.find_all('a', href=True)
            
            excel_links = []
            for link in all_links:
                href = link.get('href', '').lower()
                text = link.get_text(strip=True)
                
                # 检查链接是否指向Excel文件
                for ext in excel_extensions:
                    if ext in href or text.endswith(ext):
                        excel_links.append({
                            'url': link.get('href'),
                            'text': text if text else '未命名文件'
                        })
                        break
            
            if not excel_links:
                self.log("未发现Excel文件")
                return 0
            
            self.log(f"发现 {len(excel_links)} 个Excel文件链接")
            
            # 创建Excel文件夹
            excel_folder = os.path.join(output_folder, "excel_files")
            if not os.path.exists(excel_folder):
                os.makedirs(excel_folder)
            
            # 下载每个Excel文件
            downloaded_count = 0
            for i, excel_link in enumerate(excel_links, 1):
                if not self.is_crawling:
                    break
                    
                excel_url = excel_link['url']
                excel_text = excel_link['text']
                
                self.log(f"正在下载第 {i}/{len(excel_links)} 个Excel文件: {excel_text}")
                
                # 下载文件
                file_content = self.download_file_with_retry(excel_url, article_base_url, max_retries=2, timeout=15)
                
                if file_content:
                    # 确定文件扩展名
                    file_ext = '.xls'
                    for ext in excel_extensions:
                        if ext in excel_url.lower():
                            file_ext = ext
                            break
                    
                    # 清理文件名
                    safe_excel_text = self.clean_filename(excel_text, max_length=50)
                    if safe_excel_text == '未命名文件':
                        safe_excel_text = f"{self.clean_filename(title, max_length=30)}_附件{i}"
                    
                    # 确保文件名以扩展名结尾
                    if not safe_excel_text.endswith(file_ext):
                        safe_excel_text += file_ext
                    
                    # 保存文件
                    excel_filename = os.path.join(excel_folder, safe_excel_text)
                    
                    # 处理文件名冲突
                    counter = 1
                    original_name = excel_filename
                    while os.path.exists(excel_filename):
                        name_part, ext_part = os.path.splitext(original_name)
                        excel_filename = f"{name_part}_{counter}{ext_part}"
                        counter += 1
                    
                    try:
                        with open(excel_filename, 'wb') as f:
                            f.write(file_content)
                        
                        self.log(f"Excel文件已保存: {excel_filename}")
                        downloaded_count += 1
                        
                        # 添加延迟，避免请求过快
                        if i < len(excel_links):
                            time.sleep(1)
                    except Exception as e:
                        self.log(f"保存Excel文件失败: {e}", "WARNING")
                else:
                    self.log(f"下载Excel文件失败: {excel_url}", "WARNING")
            
            self.log(f"成功下载 {downloaded_count}/{len(excel_links)} 个Excel文件")
            return downloaded_count
            
        except Exception as e:
            self.log(f"下载Excel文件时出错: {e}", "WARNING")
            return 0

    def process_single_url(self, url, output_folder="documents"):
        """
        处理单个URL：获取内容并保存为Word文档和Excel文件
        """
        if not self.is_crawling:
            return False
            
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/142.0.0.0 Safari/537.36 Edg/142.0.0.0',
                'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8',
                'Accept-Language': 'zh-CN,zh;q=0.9,en;q=0.8'
            }
            
            self.log(f"正在处理: {url}")
            response = requests.get(url, headers=headers, timeout=self.config['timeout'])
            response.encoding = 'utf-8'
            content = response.text
            
            # 提取标题和内容
            title = self.extract_title(content)
            page_content = self.extract_content(content)
            
            if not title:
                title = "未知标题"
            
            # 清理文件名中的非法字符
            safe_title = self.clean_filename(title)
            safe_title = safe_title[:50]  # 限制文件名长度
            
            # 创建输出文件夹
            if not os.path.exists(output_folder):
                os.makedirs(output_folder)
            
            # 保存为Word文档
            doc_filename = os.path.join(output_folder, f"{safe_title}.docx")
            
            # 如果没有找到内容，尝试使用整个body
            if not page_content['html']:
                self.log("未找到特定内容区域，尝试提取整个页面内容")
                soup = BeautifulSoup(content, 'html.parser')
                body_content = soup.find('body')
                if body_content:
                    page_content = {
                        'html': body_content.prettify(),
                        'text': body_content.get_text(separator='\n', strip=True)
                    }
            
            # 如果连body都没有，使用整个页面内容
            if not page_content['html']:
                self.log("使用整个页面内容")
                page_content = {
                    'html': content,
                    'text': content
                }
            
            # 传递文章URL给save_html_to_doc函数，用于构建图片URL
            doc_success = self.save_html_to_doc(page_content['html'], doc_filename, url)
            
            # 下载Excel文件
            excel_count = self.download_excel_files(content, url, output_folder, title)
            
            if doc_success:
                self.log(f"成功处理: {title} (下载了 {excel_count} 个Excel文件)")
                return True
            else:
                self.log(f"处理失败: {title}", "WARNING")
                return False
            
        except Exception as e:
            self.log(f"处理URL {url} 时出错: {e}", "ERROR")
            return False

    def remove_duplicate_links(self, links_data):
        """
        去除重复的链接，基于URL
        """
        seen_links = set()
        unique_links = []
        
        for link_data in links_data:
            # 规范化URL，去除可能的查询参数和片段标识符
            normalized_url = link_data['link'].split('?')[0].split('#')[0]
            
            if normalized_url not in seen_links:
                seen_links.add(normalized_url)
                unique_links.append(link_data)
        
        return unique_links

    def crawl_and_process_pages(self, keyword="金融监管", start_page=1, end_page=None, output_folder="documents"):
        """
        主函数：爬取多页搜索结果并处理每个链接
        """
        if not self.is_crawling:
            return False
            
        all_links_data = []
        page_links_count = {}  # 记录每页提取的链接数量
        
        # 第一步：获取总页数
        self.log(f"开始爬取搜索结果页面，关键词: {keyword}")
        
        # 先获取第一页内容，用于确定总页数
        first_page_content = self.crawl_pbc_search(page=1, keyword=keyword)
        
        if not first_page_content:
            self.log("无法获取第一页内容，程序退出", "ERROR")
            return False
        
        # 自动获取总页数
        if end_page is None:
            total_pages = self.get_total_pages(first_page_content)
            end_page = min(total_pages, self.config['max_pages'])  # 限制最大页数
        else:
            total_pages = end_page
        
        self.log(f"计划爬取页码范围: {start_page} - {end_page}")
        
        # 第二步：爬取所有页面
        for page in range(start_page, end_page + 1):
            if not self.is_crawling:
                break
                
            self.log(f"正在爬取第 {page} 页...")
            
            # 如果是第一页，已经获取过内容，直接使用
            if page == 1:
                content = first_page_content
            else:
                content = self.crawl_pbc_search(page=page, keyword=keyword)
            
            if content:
                page_links_data = self.extract_links_with_titles_from_result_list(content)
                
                # 先去除当前页的重复链接
                unique_page_links = self.remove_duplicate_links(page_links_data)
                
                # 将当前页的后链接添加到总链接列表
                all_links_data.extend(unique_page_links)
                page_links_count[page] = len(unique_page_links)  # 记录每页后的链接数量
                
                # 立即显示后提取的链接数量
                self.log(f"第{page}页后提取到 {len(unique_page_links)} 个链接")
            
            # 添加随机延迟
            if page < end_page and self.is_crawling:
                delay = random.uniform(self.config['min_delay'], self.config['max_delay'])
                self.log(f"等待{delay:.1f}秒后获取下一页...")
                time.sleep(delay)
        
        # 去除所有页面之间的重复链接
        if self.is_crawling:
            self.log("去除所有页面间的重复链接...")
            original_count = len(all_links_data)
            all_links_data = self.remove_duplicate_links(all_links_data)
            duplicate_count = original_count - len(all_links_data)
            
            # 显示每页提取的链接数量汇总
            self.log("各页提取链接数量汇总:")
            for page, count in page_links_count.items():
                self.log(f"第{page}页: {count}个链接")
            
            self.log(f"原始链接总数: {original_count}")
            self.log(f"重复链接数: {duplicate_count}")
            self.log(f"最终后链接数: {len(all_links_data)}")
        
        # 如果后没有链接，直接返回
        if len(all_links_data) == 0:
            self.log("没有有效的链接可处理，程序结束", "WARNING")
            return False
        
        # 第三步：处理每个链接
        self.log(f"开始处理 {len(all_links_data)} 个链接...")
        success_count = 0
        
        for i, link_data in enumerate(all_links_data, 1):
            if not self.is_crawling:
                break
                
            self.log(f"[{i}/{len(all_links_data)}] 处理: {link_data['title']}")
            
            if self.process_single_url(link_data['link'], output_folder):
                success_count += 1
            
            # 处理间隔
            if i < len(all_links_data) and self.is_crawling:
                delay = random.uniform(self.config['min_delay'], self.config['max_delay'])
                self.log(f"等待{delay:.1f}秒后处理下一个链接...")
                time.sleep(delay)
        
        if self.is_crawling:
            self.log(f"处理完成: 成功 {success_count}/{len(all_links_data)} 个链接", "SUCCESS")
            self.log(f"文档保存在: {output_folder} 文件夹中")
            self.log(f"Excel文件保存在: {os.path.join(output_folder, 'excel_files')} 文件夹中")
            return success_count > 0
        else:
            self.log(f"爬取已停止: 已完成 {success_count} 个链接", "WARNING")
            return False

    def crawl_keyword(self, keyword, save_dir):
        """爬取指定关键词的文章"""
        try:
            self.is_crawling = True
            self.log(f"开始爬取中国人民银行，关键词: {keyword}")
            
            success = self.crawl_and_process_pages(keyword=keyword, output_folder=save_dir)
            
            if success:
                self.log("中国人民银行爬取任务完成", "SUCCESS")
            else:
                self.log("中国人民银行爬取任务未完成", "WARNING")
                
            return success
            
        except Exception as e:
            self.log(f"爬取过程中发生错误: {str(e)}", "ERROR")
            return False

# 保留原有的main函数用于独立运行
def main():
    import sys
    if len(sys.argv) > 1:
        keyword = sys.argv[1]
    else:
        keyword = input("请输入搜索关键词: ").strip()
    
    if not keyword:
        print("关键词不能为空")
        return
    
    crawler = PBCCrawler()
    success = crawler.crawl_keyword(keyword, f"pbc_articles_{crawler.clean_filename(keyword)}")
    
    if success:
        print("爬取任务完成")
    else:
        print("爬取任务失败")

if __name__ == "__main__":
    main()