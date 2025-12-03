from curl_cffi import requests as cffi_requests
import requests
import json
import time
import re
import os
import random
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
import urllib.parse
from PIL import Image
import io
import math

class EastMoneyCrawler:
    def __init__(self, log_callback=None):
        self.log_callback = log_callback if log_callback else print
        self.is_crawling = True  # 添加爬虫状态控制
        
        # 全局配置
        self.config = {
            'min_delay': 1,  # 最小延迟时间(秒)
            'max_delay': 3,  # 最大延迟时间(秒)
            'max_retries': 3,  # 最大重试次数
            'timeout': 15,  # 请求超时时间(秒)
        }
    
    def stop_crawling(self):
        """停止爬取"""
        self.is_crawling = False
        self.log("收到停止信号", "WARNING")
    
    def log(self, message, level="INFO"):
        """日志记录"""
        self.log_callback(f"{message}", level)
    
    def get_articles_list(self, keyword, page_index=1, page_size=10, retry_count=0):
        """获取文章列表，包含重试机制"""
        if not self.is_crawling:
            return None
            
        url = "https://search-api-web.eastmoney.com/search/jsonp"
        
        # 构建参数，包含关键词和页码
        param_data = {
            "uid": "",
            "keyword": keyword,
            "type": ["article"],
            "client": "web",
            "clientType": "web",
            "clientVersion": "curr",
            "param": {
                "article": {
                    "searchScope": "TITLE",
                    "sort": "DEFAULT",
                    "pageIndex": page_index,
                    "pageSize": page_size,
                    "preTag": "",
                    "postTag": ""
                }
            }
        }
        
        # 生成随机回调函数名和时间戳
        callback_name = f"jQuery3510{random.randint(100000000000, 999999999999)}_{int(time.time()*1000)}"
        timestamp = int(time.time()*1000) + 1
        
        params = {
            'cb': callback_name,
            'param': json.dumps(param_data),
            '_': str(timestamp)
        }
        
        headers = {
            'Accept': '*/*',
            'Accept-Language': 'zh-CN,zh;q=0.9,en;q=0.8,en-GB;q=0.7,en-US;q=0.6',
            'Cache-Control': 'no-cache',
            'Connection': 'keep-alive',
            'Cookie': 'qgqp_b_id=6df019837360b82ec24bb59d4d3deacd; st_nvi=inWJ1kAmFC6e_zWKygiqpdd21; st_si=54885234834700; nid18=038d07ffe05afe3b300e2170a3bede99; nid18_create_time=1764480048982; gviem=CU2P9X7Gpt265Got5an_Fb58e; gviem_create_time=1764480048982; fullscreengg=1; fullscreengg2=1; emshistory=%5B%22%E5%BF%AB%E6%89%8B%22%5D; st_pvi=14514299899844; st_sp=2025-11-23%2018%3A45%3A27; st_inirUrl=https%3A%2F%2Fwww.eastmoney.com%2F; st_sn=8; st_psi=20251130132119283-119101302791-7849550864; st_asi=20251130132110862-118000300908-0946980219-Web_so_ss-2',
            'Host': 'search-api-web.eastmoney.com',
            'Pragma': 'no-cache',
            'Referer': f'https://so.eastmoney.com/carticle/s?keyword={urllib.parse.quote(keyword)}',
            'Sec-Fetch-Dest': 'script',
            'Sec-Fetch-Mode': 'no-cors',
            'Sec-Fetch-Site': 'same-site',
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/142.0.0.0 Safari/537.36 Edg/142.0.0.0',
            'sec-ch-ua': '"Chromium";v="142", "Microsoft Edge";v="142", "Not_A Brand";v="99"',
            'sec-ch-ua-mobile': '?0',
            'sec-ch-ua-platform': '"Windows"'
        }
        
        try:
            # 使用curl_cffi发送请求，模拟浏览器指纹
            response = cffi_requests.get(
                url, 
                params=params, 
                headers=headers, 
                impersonate="chrome",
                timeout=self.config['timeout']
            )
            
            self.log(f"第{page_index}页请求状态码: {response.status_code}")
            
            # 处理JSONP
            content = response.text
            if content.startswith('jQuery') and content.endswith(')'):
                json_str = content[content.find('(')+1:content.rfind(')')]
                try:
                    data = json.loads(json_str)
                    return data
                except json.JSONDecodeError as e:
                    self.log(f"JSON解析错误: {e}", "WARNING")
                    return None
            else:
                try:
                    data = response.json()
                    return data
                except:
                    self.log("响应不是JSON或JSONP格式", "WARNING")
                    return None
        except Exception as e:
            self.log(f"请求文章列表时出错: {e}", "WARNING")
            
            # 重试机制
            if retry_count < self.config['max_retries']:
                retry_delay = random.uniform(2, 5)  # 重试延迟
                self.log(f"第{retry_count+1}次重试，等待{retry_delay:.1f}秒...")
                time.sleep(retry_delay)
                return self.get_articles_list(keyword, page_index, page_size, retry_count+1)
            else:
                self.log(f"已达到最大重试次数({self.config['max_retries']})，放弃请求", "ERROR")
                return None

    def get_all_articles(self, keyword):
        """获取所有页面的文章 - 解除页数限制"""
        if not self.is_crawling:
            return []
            
        all_articles = []
        
        # 先获取第一页，了解总条数和页数
        self.log(f"正在搜索关键词: {keyword}")
        first_page_data = self.get_articles_list(keyword, 1, 10)
        
        if not first_page_data:
            self.log("获取第一页数据失败", "ERROR")
            return []
        
        # 获取总条数
        total_count = first_page_data.get('hitsTotal', 0)
        self.log(f"找到 {total_count} 条相关文章")
        
        if total_count == 0:
            self.log("没有找到相关文章", "WARNING")
            return []
        
        # 计算总页数
        page_size = 10
        total_pages = math.ceil(total_count / page_size)
        self.log(f"总共 {total_pages} 页")
        
        # 解除页数限制 - 爬取所有页面
        # 添加进度提示
        if total_pages > 100:
            self.log(f"检测到大量页面({total_pages}页)，爬取可能需要较长时间...", "WARNING")
        
        max_pages = total_pages  # 爬取所有页
        
        # 添加第一页的文章
        if 'result' in first_page_data and 'article' in first_page_data['result']:
            articles = first_page_data['result']['article']
            all_articles.extend(articles)
            self.log(f"第1页: 获取到 {len(articles)} 篇文章")
        
        # 获取剩余页面的文章
        for page in range(2, total_pages + 1):
            if not self.is_crawling:
                break
                
            # 显示进度
            if page % 10 == 0 or page == total_pages:
                self.log(f"进度: {page}/{total_pages} 页 ({(page/total_pages)*100:.1f}%)")
                
            # 随机延迟，避免请求过于频繁
            delay = random.uniform(self.config['min_delay'], self.config['max_delay'])
            self.log(f"等待{delay:.1f}秒后获取第{page}页...")
            time.sleep(delay)
            
            self.log(f"正在获取第{page}页...")
            page_data = self.get_articles_list(keyword, page, page_size)
            
            if page_data and 'result' in page_data and 'article' in page_data['result']:
                articles = page_data['result']['article']
                all_articles.extend(articles)
                self.log(f"第{page}页: 获取到 {len(articles)} 篇文章")
            else:
                self.log(f"第{page}页获取失败", "WARNING")
        
        self.log(f"总共获取到 {len(all_articles)} 篇文章")
        return all_articles

    def clean_filename(self, title, max_length=100):
        """
        清理文件名，移除非法字符
        """
        if not title:
            return "untitled"
        
        # 移除HTML标签
        clean_title = re.sub(r'<[^>]+>', '', title)
        
        # 移除Windows文件名中不允许的字符
        illegal_chars = r'[<>:"/\\|?*]'
        clean_title = re.sub(illegal_chars, '', clean_title)
        
        # 移除多余的空格和换行符
        clean_title = re.sub(r'\s+', ' ', clean_title).strip()
        
        # 限制文件名长度
        if len(clean_title) > max_length:
            clean_title = clean_title[:max_length].rstrip()
        
        # 如果清理后为空，使用默认名称
        if not clean_title:
            clean_title = "untitled"
        
        return clean_title

    def extract_article_title(self, html_content):
        """从HTML内容中提取文章标题"""
        try:
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # 查找 h1 标签且 class 为 "article-title"
            title_element = soup.find('h1', class_='article-title')
            
            if title_element:
                # 提取标题文本并去除空白字符
                title_text = title_element.get_text().strip()
                return title_text
            else:
                # 如果找不到指定的标题元素，尝试其他可能的选择器
                alternative_selectors = [
                    'h1.title',
                    '.title h1',
                    'h1',
                    '.article-header h1',
                    'h1.article-header',
                    '.article-title',
                    'title'
                ]
                
                for selector in alternative_selectors:
                    title_element = soup.select_one(selector)
                    if title_element:
                        title_text = title_element.get_text().strip()
                        self.log(f"使用备用选择器找到标题: {selector}")
                        return title_text
                
                return None
                
        except Exception as e:
            self.log(f"提取标题时出错: {e}", "WARNING")
            return None

    def extract_article_content_with_images(self, html_content):
        """从HTML内容中提取文章正文和图片，保留图片位置信息"""
        try:
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # 查找正文内容
            content_div = soup.find('div', class_='xeditor_content cfh_web')
            
            if not content_div:
                # 尝试其他可能的选择器
                alternative_selectors = [
                    'div.content',
                    'div.article-content',
                    'div.main-content',
                    '.content',
                    '.article-content',
                    '.article-body',
                    '.article-text'
                ]
                
                for selector in alternative_selectors:
                    content_div = soup.select_one(selector)
                    if content_div:
                        self.log(f"使用备用选择器找到正文: {selector}")
                        break
            
            if content_div:
                # 提取所有内容，包括文本和图片
                content_elements = []
                
                # 遍历所有子元素
                for element in content_div.descendants:
                    if element.name == 'img':
                        # 处理图片元素
                        src = element.get('src')
                        if src:
                            # 处理相对URL
                            if src.startswith('//'):
                                src = 'https:' + src
                            elif src.startswith('/'):
                                src = 'https://caifuhao.eastmoney.com' + src
                            
                            content_elements.append({
                                'type': 'image',
                                'src': src,
                                'alt': element.get('alt', '')
                            })
                    elif element.name is None and element.string:
                        # 处理文本节点
                        text = element.string.strip()
                        if text:
                            content_elements.append({
                                'type': 'text',
                                'content': text
                            })
                    elif element.name in ['p', 'br']:
                        # 处理段落和换行
                        content_elements.append({
                            'type': 'newline'
                        })
                
                return content_elements
            else:
                self.log("未找到正文内容区域", "WARNING")
                return []
                
        except Exception as e:
            self.log(f"提取正文时出错: {e}", "WARNING")
            return []

    def download_image_to_memory(self, img_url, retry_count=0):
        """下载图片到内存，包含重试机制"""
        if not self.is_crawling:
            return None
            
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/142.0.0.0 Safari/537.36 Edg/142.0.0.0',
                'Referer': 'https://www.eastmoney.com/'
            }
            
            response = requests.get(img_url, headers=headers, timeout=self.config['timeout'])
            if response.status_code == 200:
                return response.content
            else:
                self.log(f"下载图片失败: {img_url} (状态码: {response.status_code})", "WARNING")
                return None
        except Exception as e:
            self.log(f"下载图片时出错: {e}", "WARNING")
            
            # 重试机制
            if retry_count < self.config['max_retries']:
                retry_delay = random.uniform(1, 3)  # 重试延迟
                self.log(f"图片下载第{retry_count+1}次重试，等待{retry_delay:.1f}秒...")
                time.sleep(retry_delay)
                return self.download_image_to_memory(img_url, retry_count+1)
            else:
                self.log(f"图片下载已达到最大重试次数({self.config['max_retries']})，放弃下载", "WARNING")
                return None

    def save_to_doc_with_images(self, article_info, content_elements, doc_save_dir):
        """将文章内容保存到Word文档，图片直接下载并插入到相应位置"""
        if not self.is_crawling:
            return None
            
        try:
            # 创建文档
            doc = Document()
            
            # 添加标题
            title = article_info.get('extracted_title') or article_info.get('list_title')
            if title:
                doc.add_heading(title, level=1)
            
            # 添加发布日期和作者
            if article_info.get('date'):
                doc.add_paragraph(f"发布日期: {article_info['date']}")
            if article_info.get('nickname'):
                doc.add_paragraph(f"作者: {article_info['nickname']}")
            
            doc.add_paragraph()  # 空行
            
            # 处理内容元素
            current_paragraph = doc.add_paragraph()
            image_index = 0
            
            for element in content_elements:
                if not self.is_crawling:
                    break
                    
                if element['type'] == 'text':
                    # 添加文本到当前段落
                    current_paragraph.add_run(element['content'] + ' ')
                elif element['type'] == 'newline':
                    # 开始新段落
                    current_paragraph = doc.add_paragraph()
                elif element['type'] == 'image':
                    # 下载图片并插入到文档
                    image_index += 1
                    self.log(f"下载并插入图片 {image_index}: {element['src']}")
                    
                    # 图片下载前随机延迟
                    img_delay = random.uniform(0.5, 1.5)
                    time.sleep(img_delay)
                    
                    img_data = self.download_image_to_memory(element['src'])
                    if img_data:
                        try:
                            # 在图片前添加换行
                            if current_paragraph.text.strip():
                                current_paragraph = doc.add_paragraph()
                            
                            # 将图片数据写入内存流
                            img_stream = io.BytesIO(img_data)
                            
                            # 添加图片到文档
                            current_paragraph.add_run().add_picture(img_stream, width=Inches(6))
                            
                            # 添加图片描述（如果有）
                            if element.get('alt'):
                                desc_paragraph = doc.add_paragraph()
                                desc_paragraph.add_run(f"图片描述: {element['alt']}").italic = True
                            
                            # 图片后添加空行
                            current_paragraph = doc.add_paragraph()
                            self.log(f"成功插入图片 {image_index}")
                        except Exception as e:
                            self.log(f"添加图片到文档时出错: {e}", "WARNING")
                            # 添加图片占位符
                            current_paragraph.add_run(f"[图片加载失败: {element.get('alt', '')}]")
                    else:
                        # 添加图片占位符
                        current_paragraph.add_run(f"[图片下载失败: {element.get('alt', '')}]")
            
            # 保存文档
            clean_title = self.clean_filename(title)
            doc_filename = f"{clean_title}.docx"
            doc_path = os.path.join(doc_save_dir, doc_filename)
            
            # 如果文件名已存在，添加序号
            counter = 1
            original_doc_path = doc_path
            while os.path.exists(doc_path):
                name, ext = os.path.splitext(original_doc_path)
                doc_path = f"{name}_{counter}{ext}"
                counter += 1
            
            doc.save(doc_path)
            return doc_path
            
        except Exception as e:
            self.log(f"保存文档时出错: {e}", "ERROR")
            return None

    def process_articles(self, articles, keyword, save_dir):
        """处理所有文章：提取内容、下载图片并保存到Word文档"""
        if not articles or not self.is_crawling:
            self.log("没有找到文章数据或爬取已停止", "WARNING")
            return []
        
        total_articles = len(articles)
        results = []
        
        # 创建保存目录
        if not os.path.exists(save_dir):
            os.makedirs(save_dir)
        
        self.log(f"开始处理 {total_articles} 篇文章...")
        
        # 创建会话，提高请求效率
        session = requests.Session()
        session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/142.0.0.0 Safari/537.36 Edg/142.0.0.0',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8',
            'Accept-Language': 'zh-CN,zh;q=0.9,en;q=0.8',
            'Referer': 'https://so.eastmoney.com/',
        })
        
        try:
            for i, article in enumerate(articles):
                if not self.is_crawling:
                    break
                    
                # 确保文章数据包含必要的字段
                if not isinstance(article, dict) or 'url' not in article or 'title' not in article:
                    self.log(f"跳过无效文章数据: {article}", "WARNING")
                    continue
                    
                url = article['url']
                list_title = article['title']
                
                # 显示处理进度
                if (i + 1) % 10 == 0 or (i + 1) == total_articles:
                    self.log(f"进度: {i+1}/{total_articles} 篇 ({(i+1)/total_articles*100:.1f}%)")
                
                self.log(f"[{i+1}/{total_articles}] 处理文章: {list_title}")
                
                try:
                    # 发送GET请求到文章URL，使用会话
                    response = session.get(url, timeout=self.config['timeout'])
                    
                    if response.status_code == 200:
                        # 提取文章标题
                        extracted_title = self.extract_article_title(response.text)
                        
                        # 提取正文内容和图片位置信息
                        content_elements = self.extract_article_content_with_images(response.text)
                        
                        # 清理标题，使其适合用作文件名
                        cleaned_title = self.clean_filename(extracted_title) if extracted_title else self.clean_filename(list_title)
                        
                        result = {
                            'index': i + 1,
                            'list_title': list_title,
                            'extracted_title': extracted_title,
                            'cleaned_title': cleaned_title,
                            'url': url,
                            'content_elements': content_elements,
                            'status_code': response.status_code,
                            'success': True
                        }
                        
                        if extracted_title:
                            self.log(f"提取的标题: {extracted_title}")
                        else:
                            self.log("未找到标题", "WARNING")
                        
                        if content_elements:
                            # 统计文本和图片数量
                            text_count = len([e for e in content_elements if e['type'] == 'text'])
                            image_count = len([e for e in content_elements if e['type'] == 'image'])
                            
                            self.log(f"找到 {text_count} 段文本和 {image_count} 张图片")
                            
                            # 保存到Word文档
                            doc_path = self.save_to_doc_with_images(result, content_elements, save_dir)
                            if doc_path:
                                result['doc_path'] = doc_path
                                self.log(f"已保存文档: {os.path.basename(doc_path)}")
                            else:
                                self.log("保存文档失败", "WARNING")
                                result['success'] = False
                        else:
                            self.log("未找到正文内容", "WARNING")
                            result['success'] = False
                            
                    else:
                        result = {
                            'index': i + 1,
                            'list_title': list_title,
                            'extracted_title': None,
                            'cleaned_title': None,
                            'url': url,
                            'content_elements': [],
                            'status_code': response.status_code,
                            'error': f'HTTP错误: {response.status_code}',
                            'success': False
                        }
                        self.log(f"HTTP错误: {response.status_code}", "WARNING")
                    
                except requests.exceptions.Timeout:
                    result = {
                        'index': i + 1,
                        'list_title': list_title,
                        'extracted_title': None,
                        'cleaned_title': None,
                        'url': url,
                        'content_elements': [],
                        'status_code': None,
                        'error': '请求超时',
                        'success': False
                    }
                    self.log("错误: 请求超时", "WARNING")
                    
                except requests.exceptions.ConnectionError:
                    result = {
                        'index': i + 1,
                        'list_title': list_title,
                        'extracted_title': None,
                        'cleaned_title': None,
                        'url': url,
                        'content_elements': [],
                        'status_code': None,
                        'error': '连接错误',
                        'success': False
                    }
                    self.log("错误: 连接错误", "WARNING")
                    
                except requests.exceptions.RequestException as e:
                    result = {
                        'index': i + 1,
                        'list_title': list_title,
                        'extracted_title': None,
                        'cleaned_title': None,
                        'url': url,
                        'content_elements': [],
                        'status_code': None,
                        'error': str(e),
                        'success': False
                    }
                    self.log(f"错误: {e}", "WARNING")
                
                results.append(result)
                
                # 随机延迟，避免请求过于频繁
                if i < total_articles - 1 and self.is_crawling:
                    delay = random.uniform(self.config['min_delay'], self.config['max_delay'])
                    self.log(f"等待{delay:.1f}秒后处理下一篇文章...")
                    time.sleep(delay)
        
        finally:
            # 确保会话关闭
            session.close()
        
        return results

    def crawl_keyword(self, keyword, save_dir):
        """爬取指定关键词的文章"""
        try:
            self.is_crawling = True
            self.log(f"开始爬取东方财富，关键词: {keyword}")
            
            # 1. 获取所有页面的文章列表
            self.log("正在获取所有文章列表...")
            all_articles = self.get_all_articles(keyword)
            
            if all_articles:
                self.log(f"成功获取 {len(all_articles)} 篇文章")
                
                # 2. 处理所有文章：提取内容、下载图片并保存到Word文档
                results = self.process_articles(all_articles, keyword, save_dir)
                
                # 3. 打印处理结果汇总
                self.print_processing_summary(results, keyword)
                
                return results
            else:
                self.log("获取文章列表失败", "WARNING")
                return []
                
        except Exception as e:
            self.log(f"爬取过程中发生错误: {str(e)}", "ERROR")
            return []

    def print_processing_summary(self, results, keyword):
        """打印处理结果汇总"""
        self.log("\n" + "="*50)
        self.log(f"关键词 '{keyword}' 的文章处理汇总")
        self.log("="*50)
        
        success_count = len([r for r in results if r.get('success', False)])
        failed_count = len(results) - success_count
        content_found = len([r for r in results if r.get('content_elements')])
        docs_saved = len([r for r in results if r.get('doc_path')])
        
        self.log(f"总文章数: {len(results)}")
        self.log(f"请求成功: {success_count}")
        self.log(f"请求失败: {failed_count}")
        self.log(f"找到正文内容: {content_found}")
        self.log(f"成功保存文档: {docs_saved}")
        
        # 显示成功保存的文档
        saved_docs = [r for r in results if r.get('doc_path')]
        if saved_docs:
            self.log(f"已保存的文档 ({len(saved_docs)}个):")
            for doc in saved_docs:
                self.log(f"  ✓ {os.path.basename(doc['doc_path'])}")
        
        # 显示失败的请求
        failed_requests = [r for r in results if not r.get('success', True)]
        if failed_requests:
            self.log(f"失败的请求 ({len(failed_requests)}个):")
            for failed in failed_requests:
                self.log(f"  ✗ {failed.get('list_title', '未知标题')}")
                self.log(f"     错误: {failed.get('error', '未知错误')}")

# 保留原有的main函数用于独立运行
def main():
    import sys
    if len(sys.argv) > 1:
        keyword = sys.argv[1]
    else:
        keyword = input("请输入搜索关键词: ").strip()
    
    if not keyword:
        print("关键词不能为空")
        return
    
    crawler = EastMoneyCrawler()
    results = crawler.crawl_keyword(keyword, f"articles_{crawler.clean_filename(keyword)}")

if __name__ == '__main__':
    main()